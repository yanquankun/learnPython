"""
@Desc: 本章节主要讲解如何文件格式转换
@Author: Mint.Yan
@Date: 2025-07-212 13:24:49
"""

from docx import Document
from openpyxl import Workbook, load_workbook
from pathlib import Path
import os

docs_path = os.path.abspath('../../files/docs/')
os.chdir(docs_path)

p = Path('.')

# 将 Word 文件转换为 Excel 文件
# 过程：
# 1. 读取 Word 文件
# 2. 创建一个新的 Excel 工作簿
# 3. 将 Word 文件中的每一段落写入 Excel 工作簿的单元格
# 4. 保存 Excel 工作簿

# 这里以 docs_path 目录下的 example.docx 文件为例

doc = Document('example.docx')

# 获取其中的所有的表格
tables = doc.tables

# data_row = []
# data_cell = []
# for t in tables:
#     for row in t.rows:
#         for cell in row.cells:
#             data_cell.append(cell.text)
#         data_row.append(data_cell)
#         data_cell = []

# 使用列表推导式简化代码
data_row = [[cell.text for cell in row.cells] for t in tables for row in t.rows]

# print(data_row)
# 可以看到如下输出：[['Qty', 'Id', 'Desc'], ['3', '101', 'Spam'], ['7', '422', 'Eggs'], ['4', '631', 'Spam, spam, eggs, and spam']]

# 创建一个新的 Excel 工作簿
wb = Workbook()
ws = wb.active
ws.title = 'Converted Data'

# 将 Word 文件中的每一段落写入 Excel 工作簿的单元格
for row in data_row:
    ws.append(row)

# 保存 Excel 工作簿
wb.save('converted.xlsx')

# 输出转换完成的提示
print("Word 文件已成功转换为 Excel 文件，保存为 'converted.xlsx'")

# 减少内存占用
# 关闭 Word 文档
doc = None
# 关闭 Excel 工作簿
wb.close()
wb = None

print("===" * 10)

# 将 Excel 文件转换为 Word 文件

# 将 docs/example.xlsx 文件转换为 Word 文件

# 读取 Excel 文件
excel_file = 'example.xlsx'
wb = load_workbook(excel_file)  # noqa
ws = wb.active

# 读取 Excel 文件中的数据
data = []

# 我们只获取有效表格数据
for row in ws.iter_rows(min_row=4, max_row=7, min_col=1, max_col=3, values_only=True):
    data.append(row)

# 创建一个新的 Word 文档
doc = Document()

# 添加表格
table = doc.add_table(rows=1, cols=len(data[0]))

# 添加表头
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(data[0]):
    hdr_cells[i].text = str(col_name)

# 添加数据行
for row in data[1:]:
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = str(cell_value)

# 保存 Word 文档
doc.save('converted.docx')

# 输出转换完成的提示
print("Excel 文件已成功转换为 Word 文件，保存为 'converted.docx'")

# 减少内存占用
doc = None
wb.close()
wb = None

# 总结
# 本章节介绍了如何将 Word 文件转换为 Excel 文件，以及如何将 Excel 文件转换为 Word 文件。
