"""
@Desc: 本章节主要讲解文件类型与对应的函数库
@Author: Mint.Yan
@Date: 2025-07-211 14:06:10
"""

# 文件类型表示文件存储的格式，Windows 根据扩展名区分文件类型
# 常⻅的音频格式：WAV、MP3
# WAV 格式没有压缩，音频无损，但文件大；MP3 格式压缩但文件小
# 常⻅的 Word 格式：docx、doc
# 常⻅的 Excel 格式：xlsx、xls

# Word 格式与 Python 函数库
# • Python 操作 Word 格式的库是 python-docx 库
# • 安装方法是 pip install python-docx
# • 导入方法是 import docx
# 官方文档：https://python-docx.readthedocs.io/en/latest/index.html

import os

# 设置files/docs/为文件所在目录
file_path = os.path.abspath('../../files/docs/')
os.chdir(file_path)

from docx import Document

# 创建一个 Word 文档
doc = Document()

# 添加标题
doc.add_heading('标题', 0)

# 加粗
p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# 标题
doc.add_heading('Heading, level 1', level=1)

# 表格
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# 增加表格
table = doc.add_table(1, 3)

# 表格首行
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty, id, desc in records:
    # 增加行
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

# 添加分页符
doc.add_page_break()

# 保存文件
doc.save('example.docx')

# Excel 格式与 Python 函数库
# • xlsx 格式一般采用 openpyxl 进行读写
# • 如果是 xls 格式，对 Word 文件读取时采用 xlrd 库，写入则需要使用 xlwt 库
# 安装方法是 pip install openpyxl
# 官方文档：https://openpyxl.readthedocs.io/en/stable/

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import datetime

# Workbook 是工作簿，可以理解为一个 Excel 文件
wb = Workbook()

# 获取当前活动的工作表，可以理解为当前打开的sheet
ws = wb.active

# 添加数据
ws.append(['123', '666'])

# 添加数据到指定单元格
ws['A2'] = 66
ws['A3'] = datetime.datetime.now()
# 设置宽度
ws.column_dimensions['A'].width = 60  # 设置A列宽度为20

# 添加多行数据
rows = [
    ['Name', 'Age', 'City'],
    ['Alice', 30, 'New York'],
    ['Bob', 25, 'Los Angeles'],
    ['Charlie', 35, 'Chicago']
]
for row in rows:
    ws.append(row)

# 设置单元格格式
ws['A1'].font = Font(name='Arial', size=14, bold=True, color='FF0000')  # 设置字体
ws['B1'].fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')  # 设置背景颜色

# 设置所有单元格的对齐方式
for row in ws.iter_rows(min_row=1, max_col=7, min_col=1, max_row=7):
    for cell in row:
        cell.alignment = Alignment(horizontal='center', vertical='center')

# 保存工作簿
wb.save('example.xlsx')
